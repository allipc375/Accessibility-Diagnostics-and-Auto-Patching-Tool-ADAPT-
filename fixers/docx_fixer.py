# docx_fixer.py


import io
import logging
import warnings

from docx import Document
from docx.shared import RGBColor

import torch
from torchvision.models import mobilenet_v3_small, MobileNet_V3_Small_Weights
from PIL import Image

logging.getLogger("transformers").setLevel(logging.ERROR)
warnings.filterwarnings("ignore", category=UserWarning)

_MN_MODEL = None
_MN_TRANSFORM = None
_MN_CATEGORIES = None


def _load_mobilenet():
    global _MN_MODEL, _MN_TRANSFORM, _MN_CATEGORIES
    if _MN_MODEL is None:
        weights = MobileNet_V3_Small_Weights.DEFAULT
        _MN_MODEL = mobilenet_v3_small(weights=weights)
        _MN_MODEL.eval()
        _MN_TRANSFORM = weights.transforms()
        _MN_CATEGORIES = weights.meta["categories"]


def _classify_image_label(image_bytes):
    """
    Returns (label, score) using mobilenet_v3_small.
    """
    _load_mobilenet()
    img = Image.open(io.BytesIO(image_bytes)).convert("RGB")
    batch = _MN_TRANSFORM(img).unsqueeze(0)

    with torch.no_grad():
        preds = _MN_MODEL(batch)
        probs = preds.softmax(dim=1)[0]
        score, idx = probs.max(dim=0)

    label = _MN_CATEGORIES[int(idx)]
    return label, float(score)


def _is_decorative_label(label: str) -> bool:
    """
    Heuristic: treat some ImageNet labels as decorative/background-like.
    """
    decorative_keywords = [
        "pattern",
        "background",
        "wallpaper",
        "border",
        "frame",
        "logo",
        "emblem",
        "symbol",
        "screen",
        "monitor",
        "web site",
        "website",
        "banner",
        "sign",
        "flag",
        "texture",
    ]
    l = label.lower()
    return any(k in l for k in decorative_keywords)


# -------------------- CONTRAST FIX --------------------

def _iter_all_paragraphs(doc: Document):
    # Body paragraphs
    for p in doc.paragraphs:
        yield p

    # Table paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def get_background_color(run):
    """
    Very rough heuristic:
    - If highlight_color is set, assume yellow background.
    - Otherwise assume white.
    (Table cell shading is not handled explicitly here.)
    """
    try:
        if run.font.highlight_color:
            return (255, 255, 0)
    except Exception:
        pass
    return (255, 255, 255)


def _is_hyperlink_run(run):
    try:
        if run.style and run.style.name == "Hyperlink":
            return True
    except Exception:
        pass
    # Fallback: check underlying rStyle
    try:
        rPr = run._r.rPr
        if rPr is not None and rPr.rStyle is not None:
            if rPr.rStyle.val == "Hyperlink":
                return True
    except Exception:
        pass
    return False


def fix_text_contrast(doc: Document, issues, min_ratio=4.5):
    """
    Fixes low contrast text in:
    - Body paragraphs
    - Table cell paragraphs
    - Hyperlinks (even if they had no explicit color)
    """
    fixed = 0

    for p in _iter_all_paragraphs(doc):
        for run in p.runs:
            try:
                bg = get_background_color(run)

                # Hyperlinks: if no explicit color, set to accessible dark blue
                if _is_hyperlink_run(run) and (
                    run.font.color is None or run.font.color.rgb is None
                ):
                    run.font.color.rgb = RGBColor(0, 0, 139)  # dark blue
                    run.font.underline = True
                    issues.append(
                        ("Hyperlink color set for contrast", p.text[:50] or "Hyperlink")
                    )
                    fixed += 1
                    continue

                if run.font.color is not None and run.font.color.rgb is not None:
                    rgb = tuple(run.font.color.rgb)
                    if AllyScores.contrast_ratio(rgb, bg) < min_ratio:
                        black_ratio = AllyScores.contrast_ratio((0, 0, 0), bg)
                        white_ratio = AllyScores.contrast_ratio((255, 255, 255), bg)

                        if black_ratio >= white_ratio:
                            run.font.color.rgb = RGBColor(0, 0, 0)
                        else:
                            run.font.color.rgb = RGBColor(255, 255, 255)

                        issues.append(
                            ("Low contrast text fixed", p.text[:50] or "Paragraph")
                        )
                        fixed += 1
            except Exception:
                continue

    return fixed


# -------------------- LANGUAGE FIX --------------------

def fix_language(doc: Document, issues):
    if not doc.core_properties.language:
        doc.core_properties.language = "en-US"
        issues.append(("Language set to en-US", "Document"))
        return 1
    return 0


# -------------------- TABLE HEADER FIX --------------------

def fix_table_headers(doc: Document, issues):
    fixed = 0

    for i, table in enumerate(doc.tables):
        try:
            first_row = table.rows[0]._tr
            trPr = first_row.get_or_add_trPr()

            headers = trPr.xpath("./w:tblHeader")
            if not headers:
                from docx.oxml import OxmlElement
                header = OxmlElement("w:tblHeader")
                trPr.append(header)
                issues.append(("Table header added", f"Table {i+1}"))
                fixed += 1
        except Exception:
            continue

    return fixed


# -------------------- ALT TEXT & DECORATIVE FIX --------------------

def _get_image_bytes_from_shape(doc: Document, shape):
    """
    Extract raw image bytes from an inline shape.
    """
    try:
        inline = shape._inline
        blip = inline.graphic.graphicData.pic.blipFill.blip
        embed_id = blip.embed
        image_part = doc.part.related_parts[embed_id]
        return image_part.blob
    except Exception:
        return None


def fix_alt_text_and_decorative(doc: Document, issues):
    """
    Uses mobilenet_v3_small to:
    - Classify each image.
    - Mark likely decorative images as decorative:
        title="Decorative", descr="".
    - For informative images with missing alt text, add:
        descr="Image of <label>".
    - Normalize existing decorative images so they don't expose alt text.
    """
    fixed_missing_alt = 0
    fixed_decorative = 0

    for i, shape in enumerate(doc.inline_shapes):
        try:
            docPr = shape._inline.docPr
            descr = docPr.get("descr")
            title = docPr.get("title")

            # Existing decorative marking
            is_marked_decorative = (
                (title and title.strip().lower() == "decorative")
                or (descr and descr.strip().lower() == "decorative")
            )

            image_bytes = _get_image_bytes_from_shape(doc, shape)
            label = None
            score = None

            if image_bytes is not None:
                try:
                    label, score = _classify_image_label(image_bytes)
                except Exception:
                    label, score = None, None

            # Decide if the image is decorative based on label
            is_pred_decorative = False
            if label is not None:
                is_pred_decorative = _is_decorative_label(label)

            # 1. Normalize already-marked decorative images
            if is_marked_decorative or is_pred_decorative:
                changed = False

                # Clear alt text for decorative images
                if descr not in (None, ""):
                    docPr.set("descr", "")
                    changed = True

                # Normalize title
                if title != "Decorative":
                    docPr.set("title", "Decorative")
                    changed = True

                if changed:
                    issues.append(
                        ("Decorative image normalized", f"Image {i+1}")
                    )
                    fixed_decorative += 1
                continue

            # 2. Informative images: ensure alt text
            has_alt = bool(descr) or bool(title)
            if not has_alt:
                if label is not None:
                    alt_text = f"Image of {label}"
                else:
                    alt_text = "Image"

                docPr.set("descr", alt_text)
                issues.append(
                    ("Alt text added to image", f"Image {i+1}")
                )
                fixed_missing_alt += 1

        except Exception:
            continue

    return fixed_missing_alt, fixed_decorative


# -------------------- SCORING --------------------

class AllyScores:

    @staticmethod
    def contrast_ratio(c1, c2):
        def lum(c):
            c = [x / 255.0 for x in c]

            def chan(v):
                if v <= 0.03928:
                    return v / 12.92
                return ((v + 0.055) / 1.055) ** 2.4

            return (
                0.2126 * chan(c[0])
                + 0.7152 * chan(c[1])
                + 0.0722 * chan(c[2])
            )

        L1, L2 = lum(c1), lum(c2)
        return (max(L1, L2) + 0.05) / (min(L1, L2) + 0.05)

    @staticmethod
    def score_from_count(count):
        if count == 0:
            return 100
        if count <= 5:
            return 76
        if count <= 21:
            return 53
        return 5

    @staticmethod
    def compute(
        missing_alt,
        decorative,
        language_missing,
        headings_missing,
        tables_missing,
        contrast,
        lists,
        links,
    ):
        scores = {
            "alternative_text": AllyScores.score_from_count(missing_alt),
            "decorative": AllyScores.score_from_count(decorative),
            "language": 95 if language_missing else 100,
            "headings": 99 if headings_missing else 100,
            "tables": 68 if tables_missing else 100,
            "color_contrast": AllyScores.score_from_count(contrast),
            "lists": AllyScores.score_from_count(lists),
            "links": AllyScores.score_from_count(links),
        }

        final = min(scores.values())

        if final == 100:
            band = "Dark Green"
        elif final >= 67:
            band = "Light Green"
        elif final >= 34:
            band = "Yellow"
        else:
            band = "Red"

        scores["final"] = final
        scores["band"] = band
        return scores
