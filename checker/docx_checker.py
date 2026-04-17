# docx_checker.py

from docx import Document
from fixers.docx_fixer import (
    fix_text_contrast,
    fix_language,
    fix_table_headers,
    fix_alt_text_and_decorative,
    AllyScores,
)


def _iter_all_paragraphs(doc):
    # Body paragraphs
    for p in doc.paragraphs:
        yield p

    # Table paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def _check_images_alt_and_decorative(doc, issues):
    """
    Returns (missing_alt_count, decorative_issue_count).
    - missing_alt_count: images with no alt text and not marked decorative
    - decorative_issue_count: images marked decorative but still have alt text or inconsistent marking
    """
    missing_alt = 0
    decorative_issues = 0

    for i, shape in enumerate(doc.inline_shapes):
        try:
            docPr = shape._inline.docPr
            descr = docPr.get("descr")
            title = docPr.get("title")

            is_decorative = (
                (title and title.strip().lower() == "decorative")
                or (descr and descr.strip().lower() == "decorative")
            )

            if is_decorative:
                # Decorative images should not expose meaningful alt text
                if descr not in (None, "", "decorative"):
                    decorative_issues += 1
                    issues.append(
                        ("Decorative image has alt text", f"Image {i+1}")
                    )
            else:
                # Not decorative: must have some alt text (title or descr)
                if not descr and not title:
                    missing_alt += 1
                    issues.append(
                        ("Image missing alt text", f"Image {i+1}")
                    )
        except Exception:
            continue

    return missing_alt, decorative_issues


def check_docx(file_path, apply_fix=False):
    print(f"Running DOCX accessibility check for: {file_path}")
    doc = Document(file_path)

    issues = []

    # -------- LANGUAGE --------
    language_missing = 0 if doc.core_properties.language else 1
    if language_missing:
        issues.append(("Document language not set", "Document"))

    # -------- TABLES --------
    tables_missing = 0
    for i, table in enumerate(doc.tables):
        try:
            first_row = table.rows[0]._tr
            trPr = first_row.xpath("./w:trPr")
            header = len(trPr[0].xpath("./w:tblHeader")) if trPr else 0
            if header == 0:
                tables_missing += 1
                issues.append(("Table missing header row", f"Table {i+1}"))
        except Exception:
            tables_missing += 1
            issues.append(("Table header could not be determined", f"Table {i+1}"))

    # -------- CONTRAST (BODY, TABLES, LINKS) --------
    contrast = 0
    for p in _iter_all_paragraphs(doc):
        for run in p.runs:
            try:
                # We only check contrast when an explicit RGB color is set
                if run.font.color is not None and run.font.color.rgb is not None:
                    rgb = tuple(run.font.color.rgb)
                    # Assume white background for scoring; fixer uses a better heuristic
                    if AllyScores.contrast_ratio(rgb, (255, 255, 255)) < 4.5:
                        contrast += 1
                        location = p.text[:50] or "Paragraph"
                        issues.append(("Low contrast text", location))
            except Exception:
                continue

    # -------- IMAGES: ALT TEXT & DECORATIVE --------
    missing_alt, decorative_issues = _check_images_alt_and_decorative(doc, issues)

    # -------- APPLY FIX --------
    if apply_fix:
        fix_text_contrast(doc, issues)
        fix_language(doc, issues)
        fix_table_headers(doc, issues)
        fix_alt_text_and_decorative(doc, issues)

        fixed_path = file_path.replace(".docx", "_fixed.docx")
        doc.save(fixed_path)
        print(f"Saved fixed file: {fixed_path}")

    # -------- SCORING --------
    scores = AllyScores.compute(
        missing_alt=missing_alt,
        decorative=decorative_issues,
        language_missing=language_missing,
        headings_missing=0,
        tables_missing=tables_missing,
        contrast=contrast,
        lists=0,
        links=0,
    )

    return {
        "score": scores["final"],
        "band": scores["band"],
        "details": scores,
        "issues": issues,
    }
